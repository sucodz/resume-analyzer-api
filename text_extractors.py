# text_extractors.py
import docx
import PyPDF2

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    return " ".join(page.extract_text() for page in reader.pages if page.extract_text())

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    return "\n".join(p.text for p in doc.paragraphs)









# def extract_text_from_pdf(pdf_file):
#     import PyPDF2
#     if isinstance(pdf_file, str):
#         with open(pdf_file, "rb") as f:
#             reader = PyPDF2.PdfReader(f)
#             return " ".join(page.extract_text() for page in reader.pages if page.extract_text())
#     else:
#         reader = PyPDF2.PdfReader(pdf_file)
#         return " ".join(page.extract_text() for page in reader.pages if page.extract_text())

# def extract_text_from_docx(docx_file):
#     import docx
#     if isinstance(docx_file, str):
#         doc = docx.Document(docx_file)
#     else:
#         doc = docx.Document(docx_file)
#     return "\n".join(p.text for p in doc.paragraphs)








# # text_extractors.py
# import docx
# import PyPDF2

# def extract_text_from_pdf(pdf_file):
#     reader = PyPDF2.PdfReader(pdf_file)
#     return " ".join(page.extract_text() for page in reader.pages if page.extract_text())

# def extract_text_from_docx(docx_file):
#     doc = docx.Document(docx_file)
#     return "\n".join([p.text for p in doc.paragraphs])
